import customtkinter as ctk
import subprocess
import threading
from pathlib import Path
from tkinter import filedialog, messagebox

# Third-party libs
from pdf2docx import Converter
from docx import Document
from fpdf import FPDF

class DocConverter(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # UI Config
        self.title("OpenConv - MultiFormat")
        self.geometry("600x480")
        ctk.set_appearance_mode("dark")
        
        self.src_path = None
        self._init_ui()

    def _init_ui(self):
        """Minimalist interface setup"""
        self.main_lbl = ctk.CTkLabel(self, text="Document Engine", font=("Segoe UI", 20, "bold"))
        self.main_lbl.pack(pady=(30, 10))

        self.btn_browse = ctk.CTkButton(self, text="Select Source File", 
                                       fg_color="#34495e", hover_color="#2c3e50",
                                       command=self._pick_file)
        self.btn_browse.pack(pady=10)

        self.info_lbl = ctk.CTkLabel(self, text="Wait for file selection...", font=("Arial", 11), text_color="gray")
        self.info_lbl.pack(pady=5)

        # Mode Selection
        self.ops = ["Word -> PDF", "PDF -> Word", "TXT -> PDF", "TXT -> Word", "Word -> TXT"]
        self.mode_sel = ctk.CTkOptionMenu(self, values=self.ops, width=200)
        self.mode_sel.set("Word -> PDF")
        self.mode_sel.pack(pady=25)

        # Execute
        self.exec_btn = ctk.CTkButton(self, text="START CONVERSION", state="disabled",
                                     height=40, font=("Arial", 13, "bold"),
                                     command=self._run_thread)
        self.exec_btn.pack(pady=20)

    def _pick_file(self):
        path = filedialog.askopenfilename()
        if path:
            self.src_path = Path(path)
            self.info_lbl.configure(text=self.src_path.name, text_color="white")
            self.exec_btn.configure(state="normal")

    def _run_thread(self):
        threading.Thread(target=self._engine, daemon=True).start()

    def _engine(self):
        p = self.src_path
        mode = self.mode_sel.get()
        out_dir = p.parent
        target_name = p.stem
        
        self.exec_btn.configure(state="disabled", text="Running...")
        
        try:
            if mode == "Word -> PDF":
                # Using LibreOffice for professional Linux-based conversion
                subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', str(p), '--outdir', str(out_dir)], check=True)
            
            elif mode == "PDF -> Word":
                cv = Converter(str(p))
                cv.convert(str(out_dir / f"{target_name}.docx"))
                cv.close()

            elif mode == "TXT -> PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=11)
                with open(p, "r", encoding="utf-8") as f:
                    for line in f:
                        pdf.cell(0, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                pdf.output(str(out_dir / f"{target_name}.pdf"))

            elif mode == "TXT -> Word":
                doc = Document()
                with open(p, "r", encoding="utf-8") as f:
                    doc.add_paragraph(f.read())
                doc.save(str(out_dir / f"{target_name}.docx"))

            elif mode == "Word -> TXT":
                doc = Document(p)
                content = "\n".join([para.text for para in doc.paragraphs])
                with open(str(out_dir / f"{target_name}.txt"), "w", encoding="utf-8") as f:
                    f.write(content)

            messagebox.showinfo("Success", "Process finished.")
        except Exception as e:
            messagebox.showerror("Error", f"Details: {str(e)}")
        finally:
            self.exec_btn.configure(state="normal", text="START CONVERSION")

if __name__ == "__main__":
    app = DocConverter()
    app.mainloop()
